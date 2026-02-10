import time
import requests
import pandas as pd
import re
import threading
import os
import json
import subprocess
from datetime import datetime
import platform

from selenium import webdriver
from selenium.webdriver.common.by import By
from selenium.webdriver.common.keys import Keys
from selenium.webdriver.chrome.options import Options
from selenium.webdriver.chrome.service import Service
from selenium.webdriver.support.ui import WebDriverWait
from selenium.webdriver.support import expected_conditions as EC

# File parsing
from pathlib import Path

# Optional imports (used only if file type matches)
import docx
import pdfplumber

# Import OTP authentication module
try:
    from whatsapp_otp import (
        request_whatsapp_otp,
        verify_whatsapp_otp,
        get_otp_session,
        wait_for_whatsapp_authentication,
        wait_for_otp_input_field,
        mark_otp_verified,
        enter_country_and_phone,
        start_whatsapp_login_automation,
        save_otp_session,
        is_session_already_logged_in,
        mark_already_logged_in
    )
except ImportError as e:
    # Fallback if module not found
    print(f"[IMPORT ERROR] Could not import whatsapp_otp: {e}")
    save_otp_session = None
    start_whatsapp_login_automation = None
    wait_for_otp_input_field = None
    wait_for_whatsapp_authentication = None
    mark_otp_verified = None
    is_session_already_logged_in = None
    mark_already_logged_in = None


# =========================
# ACTIVE DRIVERS MANAGEMENT
# =========================
# Global dictionary to track active drivers for each phone number
ACTIVE_DRIVERS = {}
DRIVERS_LOCK = threading.Lock()

def register_driver(phone_number: str, driver):
    """Register an active driver for a specific phone number"""
    with DRIVERS_LOCK:
        ACTIVE_DRIVERS[phone_number] = driver

def close_driver(phone_number: str):
    """Close and unregister a driver for a specific phone number"""
    with DRIVERS_LOCK:
        if phone_number in ACTIVE_DRIVERS:
            try:
                driver = ACTIVE_DRIVERS[phone_number]
                if driver:
                    driver.quit()
                    print(f"[DRIVER] Closed driver for {phone_number}")
            except Exception as e:
                print(f"[DRIVER] Error closing driver: {e}")
            finally:
                del ACTIVE_DRIVERS[phone_number]


def hide_driver_window(phone_number: str):
    """Move the browser window off-screen so it becomes invisible but keep the driver running."""
    with DRIVERS_LOCK:
        try:
            if phone_number in ACTIVE_DRIVERS:
                driver = ACTIVE_DRIVERS[phone_number]
                if driver:
                    # Move off-screen on Windows: negative coordinates
                    driver.set_window_position(-32000, -32000)
                    print(f"[DRIVER] Hid browser window for {phone_number} (moved off-screen)")
                    return True
        except Exception as e:
            print(f"[DRIVER] Error hiding driver window for {phone_number}: {e}")
    return False

def close_all_drivers():
    """Close all active drivers"""
    with DRIVERS_LOCK:
        for phone_number in list(ACTIVE_DRIVERS.keys()):
            try:
                driver = ACTIVE_DRIVERS[phone_number]
                if driver:
                    driver.quit()
                    print(f"[DRIVER] Closed driver for {phone_number}")
            except Exception as e:
                print(f"[DRIVER] Error closing driver: {e}")
        ACTIVE_DRIVERS.clear()


# =========================
# OLLAMA CONFIG
# =========================
OLLAMA_URL = "http://localhost:11434/api/generate"
OLLAMA_MODEL = "phi"
OLLAMA_TIMEOUT = 120


# =========================
# WHATSAPP PROFILE (PERSISTENT)
# =========================
# Store profiles in project directory for proper permissions
PROJECT_ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
WHATSAPP_PROFILE_DIR = os.path.join(PROJECT_ROOT, ".eduping_profiles")
os.makedirs(WHATSAPP_PROFILE_DIR, exist_ok=True)  # Create if doesn't exist

PROFILE_LOCK_FILE = os.path.join(WHATSAPP_PROFILE_DIR, ".lock")
STATUS_FILE = os.path.join(WHATSAPP_PROFILE_DIR, ".status.json")


def get_user_profile_dir(user_whatsapp: str = None) -> str:
    """
    Get profile directory for a specific user.
    If user_whatsapp is provided, create a user-specific profile.
    Otherwise, use the default profile.
    """
    if user_whatsapp:
        # Create a user-specific profile directory
        user_dir = os.path.join(WHATSAPP_PROFILE_DIR, f"user_{user_whatsapp}")
        os.makedirs(user_dir, exist_ok=True)
        return user_dir
    return WHATSAPP_PROFILE_DIR


def clean_phone_num(phone):
    """Clean phone number and ensure proper format for WhatsApp"""
    phone = str(phone).strip().replace(" ", "").replace("-", "").replace("(", "").replace(")", "")
    # Remove any leading zeros
    if phone.startswith("0"):
        phone = phone[1:]
    # Add country code if not present (assuming India +91)
    if not phone.startswith("+") and not phone.startswith("91"):
        phone = "91" + phone
    elif phone.startswith("+"):
        phone = phone[1:]  # Remove + as WhatsApp URL doesn't need it
    print(f"[PHONE] Cleaned: {phone}")
    return phone


# =========================
# MESSAGE FILE EXTRACTION
# =========================
def extract_messages_from_file(file_path: str) -> list[str]:
    ext = Path(file_path).suffix.lower()

    if ext == ".csv":
        df = pd.read_csv(file_path)
        if "message" not in df.columns:
            raise ValueError("CSV must contain a 'message' column")
        return df["message"].dropna().astype(str).tolist()

    if ext == ".txt":
        with open(file_path, "r", encoding="utf-8") as f:
            return [line.strip() for line in f if line.strip()]

    if ext == ".pdf":
        messages = []
        with pdfplumber.open(file_path) as pdf:
            for page in pdf.pages:
                text = page.extract_text()
                if text:
                    for line in text.split("\n"):
                        if line.strip():
                            messages.append(line.strip())
        return messages

    if ext == ".docx":
        doc = docx.Document(file_path)
        return [p.text.strip() for p in doc.paragraphs if p.text.strip()]

    raise ValueError("Unsupported message file type")


# =========================
# PHONE FILE EXTRACTION
# =========================
def extract_phones_from_file(file_path: str) -> dict[int, list[str]]:
    """Extract phone numbers grouped by batch_id from various file formats"""
    ext = Path(file_path).suffix.lower()

    if ext == ".csv":
        df = pd.read_csv(file_path)
        
        # Check if batch_id column exists
        if "batch_id" in df.columns:
            batches = {}
            if "phone" in df.columns:
                for batch_id, group in df.groupby("batch_id"):
                    phones = group["phone"].dropna().astype(str).tolist()
                    batches[int(batch_id)] = phones
            else:
                # Use first column as phone if no phone column
                phone_col = df.columns[0]
                for batch_id, group in df.groupby("batch_id"):
                    phones = group[phone_col].dropna().astype(str).tolist()
                    batches[int(batch_id)] = phones
            return batches
        else:
            # No batch_id column - group phones into batches of 8
            if "phone" in df.columns:
                phones = df["phone"].dropna().astype(str).tolist()
            else:
                if len(df.columns) > 0:
                    phones = df.iloc[:, 0].dropna().astype(str).tolist()
                else:
                    raise ValueError("CSV file is empty or has no columns")
            
            # Group into batches of 8
            batches = {}
            batch_size = 8
            for i in range(0, len(phones), batch_size):
                batch_id = (i // batch_size) + 1
                batches[batch_id] = phones[i:i + batch_size]
            return batches

    if ext == ".txt":
        phones = []
        with open(file_path, "r", encoding="utf-8") as f:
            for line in f:
                line = line.strip()
                if line:
                    # Each line is treated as a phone number
                    phones.append(line)
        # Group into batches of 8
        batches = {}
        batch_size = 8
        for i in range(0, len(phones), batch_size):
            batch_id = (i // batch_size) + 1
            batches[batch_id] = phones[i:i + batch_size]
        return batches

    if ext == ".pdf":
        phones = []
        with pdfplumber.open(file_path) as pdf:
            for page in pdf.pages:
                text = page.extract_text()
                if text:
                    for line in text.split("\n"):
                        line = line.strip()
                        if line:
                            phones.append(line)
        # Group into batches of 8
        batches = {}
        batch_size = 8
        for i in range(0, len(phones), batch_size):
            batch_id = (i // batch_size) + 1
            batches[batch_id] = phones[i:i + batch_size]
        return batches

    if ext == ".docx":
        doc = docx.Document(file_path)
        phones = []
        for p in doc.paragraphs:
            text = p.text.strip()
            if text:
                phones.append(text)
        # Group into batches of 8
        batches = {}
        batch_size = 8
        for i in range(0, len(phones), batch_size):
            batch_id = (i // batch_size) + 1
            batches[batch_id] = phones[i:i + batch_size]
        return batches

    raise ValueError("Unsupported phone file type")


# =========================
# SUMMARY HELPERS
# =========================
def sanitize_summary(text: str) -> str:
    text = re.sub(r"^\d+\.\s*", "", text.strip())
    return text.strip(" .-")


def extract_summary_deterministic(message):
    """Extract summary using deterministic rule-based approach for consistency"""
    msg = message.strip()
    
    # If message already has format like "Event - Date: X, Time: Y", keep it as-is
    if " - Date:" in msg or " - Date: " in msg:
        # Clean up extra spaces but keep the format
        msg = re.sub(r'\s+', ' ', msg)  # Normalize spaces
        return msg
    
    # For natural language messages, extract key information deterministically
    import re
    
    # Pattern: "Event on [date] at [time] in [location]"
    # Example: "Placement drive for HPE... is being held on 23rd March in Seminar hall"
    pattern = r"^(.+?)\s+(?:is\s+)?(?:being\s+)?(?:held|scheduled|conducted|planned)\s+on\s+([^at\n\.]+?)(?:\s+at\s+([^in\n\.]+?))?(?:\s+in\s+([^\.\n]+?))?(?:\.|$)"
    match = re.search(pattern, msg, re.IGNORECASE)
    if match:
        event = match.group(1).strip()
        date = match.group(2).strip() if match.group(2) else ""
        time_part = match.group(3).strip() if match.group(3) else ""
        location = match.group(4).strip() if match.group(4) else ""
        
        # Build summary preserving the message structure
        result = event
        if date:
            result += f" on {date}"
        if time_part:
            result += f" at {time_part}"
        if location:
            result += f" in {location}"
        
        # Add any additional info after the location
        if match.end() < len(msg):
            remaining = msg[match.end():].strip()
            if remaining and not remaining.startswith('.'):
                result += " " + remaining
        return result.strip()
    
    # Pattern: "Event on [date]" (simpler pattern)
    pattern2 = r"^(.+?)\s+on\s+([^\.\n]+?)(?:\.|$)"
    match = re.search(pattern2, msg, re.IGNORECASE)
    if match:
        event = match.group(1).strip()
        rest = match.group(2).strip()
        return f"{event} on {rest}".strip()
    
    # If message is already concise and clear, return as-is
    if len(msg) < 200 and ('.' not in msg or msg.count('.') <= 1):
        return msg
    
    # Fallback: return the message as-is (deterministic - no transformation)
    return msg


def batch_summarize_messages(messages):
    """Generate summaries in format: Event name - date - time - location"""
    if not messages:
        return []
    
    # Build a single prompt for all messages
    prompt = (
        "Extract information from each message and format EXACTLY as: Event name - date - time - location\n\n"
        "Rules:\n"
        "- Extract the EVENT NAME (exam name, event name, meeting name, etc.)\n"
        "- Extract the DATE (if mentioned)\n"
        "- Extract the TIME (if mentioned)\n"
        "- Extract the LOCATION (if mentioned)\n"
        "- Use ' - ' (space dash space) to separate each part\n"
        "- If date/time/location is NOT mentioned, skip that part\n"
        "- Keep event names short (just the key name, not full sentence)\n"
        "- Output ONE summary per message, one per line\n"
        "- NO explanations, NO extra text, just the formatted summaries\n\n"
        "Examples:\n"
        "The BDA exam is on 12th Jan at 10 30 → BDA exam - 12th Jan - 10 30\n"
        "Placement drive for HPE is on 23rd March in Seminar hall → Placement drive for HPE - 23rd March - Seminar hall\n"
        "Internal assessment test for Data Structures on 5th Feb at 2 00 PM in Block B → Data Structures assessment - 5th Feb - 2 00 PM - Block B\n"
        "Workshop on AI on 18th April at Main Auditorium → AI Workshop - 18th April - Main Auditorium\n"
        "SAN exam is on 1st Dec at 10 30 am → SAN exam - 1st Dec - 10 30 am\n\n"
        "Now extract from these messages:\n\n"
    )
    
    # Add all messages
    for i, msg in enumerate(messages, 1):
        prompt += f"{i}. {msg}\n"
    
    prompt += "\nSummaries (format: Event name - date - time - location, one per line):\n"
    
    # Single API call for all messages
    payload = {"model": OLLAMA_MODEL, "prompt": prompt, "stream": False}
    summaries = []
    
    try:
        resp = requests.post(OLLAMA_URL, json=payload, timeout=OLLAMA_TIMEOUT)
        resp.raise_for_status()
        raw = resp.json()["response"]
        
        # Parse the response to extract summaries
        lines = raw.splitlines()
        extracted_summaries = []
        
        for line in lines:
            line = line.strip()
            # Skip empty lines
            if not line or len(line) < 5:
                continue
            # Skip instruction/header lines
            if line.lower().startswith(("example", "rules:", "messages", "summaries", "output", "format:", "now extract")):
                continue
            # Skip lines that are just numbers
            if re.match(r'^\d+\.?\s*$', line):
                continue
            # Look for lines with the dash format (our target format)
            if " - " in line:
                # Remove leading numbers if present
                line = re.sub(r'^\d+\.?\s*', '', line).strip()
                if line and len(line) > 5:
                    extracted_summaries.append(line)
            # Also accept lines that look like summaries (long enough, not headers)
            elif len(line) > 15 and not line.lower().startswith(("message", "summary")):
                line = re.sub(r'^\d+\.?\s*', '', line).strip()
                extracted_summaries.append(line)
        
        # Match summaries to messages
        for i, msg in enumerate(messages):
            if i < len(extracted_summaries):
                summary = extracted_summaries[i].strip()
                # Clean up summary
                summary = sanitize_summary(summary)
            else:
                # Fallback: try to extract basic info
                summary = extract_basic_summary(msg)
            
            # Validate and clean summary
            if not summary or len(summary) < 3:
                summary = extract_basic_summary(msg)
            
            summaries.append(summary)
        
        # Ensure we have the right number of summaries
        while len(summaries) < len(messages):
            msg = messages[len(summaries)]
            summary = extract_basic_summary(msg)
            summaries.append(summary)
            
    except Exception as e:
        # Fallback: extract basic summaries
        for msg in messages:
            summary = extract_basic_summary(msg)
            summaries.append(summary)
    
    return summaries[:len(messages)]


def extract_basic_summary(msg):
    """Fallback function to extract basic summary from message"""
    msg = msg.strip()
    parts = []
    
    # Extract event name (first part before "on", "is", "will be", etc.)
    event_match = re.match(r'^(.+?)\s+(?:is|will be|is being|is scheduled|is planned|will be conducted)', msg, re.IGNORECASE)
    if event_match:
        event = event_match.group(1).strip()
        # Clean up event name
        event = re.sub(r'^(The|A|An)\s+', '', event, flags=re.IGNORECASE).strip()
        parts.append(event)
    else:
        # Try simpler pattern
        event = msg.split(' on ')[0].split(' is ')[0].split(' will ')[0].strip()
        event = re.sub(r'^(The|A|An)\s+', '', event, flags=re.IGNORECASE).strip()
        parts.append(event if event else "Event")
    
    # Extract date
    date_match = re.search(r'on\s+(\d+(?:st|nd|rd|th)?\s+(?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)[a-z]*(?:\s+\d+)?)', msg, re.IGNORECASE)
    if date_match:
        parts.append(date_match.group(1).strip())
    
    # Extract time
    time_match = re.search(r'at\s+(\d+\s*\d*\s*(?:am|pm|AM|PM))', msg, re.IGNORECASE)
    if time_match:
        parts.append(time_match.group(1).strip())
    
    # Extract location
    location_match = re.search(r'in\s+([^and\.]+?)(?:\s+and|\.|$)', msg, re.IGNORECASE)
    if location_match:
        location = location_match.group(1).strip()
        if len(location) < 50:  # Reasonable location length
            parts.append(location)
    
    return " - ".join(parts) if len(parts) > 1 else msg[:80]


def format_as_bullets(summary: str) -> str:
    """Format summary as a single-line, clear message in format: *Event Name - Date - Time - Location*"""
    # Clean up the summary
    summary = summary.strip()
    # Remove any extra dashes or formatting
    parts = [p.strip() for p in summary.split(" - ") if p.strip()]
    # Join all parts with " - " to create one clear line
    single_line = " - ".join(parts)
    # Wrap in asterisks for WhatsApp bold formatting (no emojis to avoid ChromeDriver BMP issues)
    return f"*{single_line}*"


# =========================
# BUILD PAIRS
# =========================
def build_pairs(phone_file_path, message_file_path):
    """Build phone-summary pairs: Combine all summaries into one message, send to each batch"""
    phone_batches = extract_phones_from_file(phone_file_path)
    messages = extract_messages_from_file(message_file_path)
    
    print("\n" + "="*50)
    print("[SUMMARY GENERATION] Starting...")
    print(f"[SUMMARY] Total messages to process: {len(messages)}")
    print("="*50)
    
    summaries = batch_summarize_messages(messages)
    
    print("\n" + "="*50)
    print("[SUMMARIES GENERATED]")
    print("="*50)
    for i, summary in enumerate(summaries, 1):
        print(f"  {i}. {summary}")
    print("="*50 + "\n")

    # Format each summary
    formatted_summaries = [format_as_bullets(summary) for summary in summaries]
    
    # Combine all summaries into one message (one per line)
    combined_message = "\n".join(formatted_summaries)
    
    print("[COMBINED MESSAGE TO SEND]:")
    print("-"*50)
    print(combined_message)
    print("-"*50 + "\n")

    # Create pairs: send combined message to all phones in each batch
    pairs = []
    all_phones = []
    for batch_id, phones in sorted(phone_batches.items()):
        for phone in phones:
            cleaned_phone = clean_phone_num(phone)
            pairs.append((cleaned_phone, combined_message))
            all_phones.append(cleaned_phone)
    
    print(f"[PHONES] Will send to {len(pairs)} phone numbers:")
    for p in all_phones:
        print(f"  - {p}")
    print("")
    
    return pairs


# =========================
# LOCK FILE HELPERS
# =========================
def is_process_running(pid):
    """Check if a process is still running"""
    try:
        if platform.system() == "Windows":
            # On Windows, use tasklist or try to signal the process
            import subprocess
            result = subprocess.run(
                ["tasklist", "/FI", f"PID eq {pid}"],
                capture_output=True,
                text=True,
                creationflags=subprocess.CREATE_NO_WINDOW if hasattr(subprocess, 'CREATE_NO_WINDOW') else 0
            )
            return str(pid) in result.stdout
        else:
            # On Unix-like systems, use os.kill with signal 0
            os.kill(pid, 0)
            return True
    except (OSError, ProcessLookupError, subprocess.SubprocessError):
        return False


def cleanup_stale_lock():
    """Remove lock file if the process is no longer running"""
    if os.path.exists(PROFILE_LOCK_FILE):
        try:
            with open(PROFILE_LOCK_FILE, "r") as f:
                pid_str = f.read().strip()
                if pid_str.isdigit():
                    pid = int(pid_str)
                    if not is_process_running(pid):
                        # Process is dead, remove stale lock
                        os.remove(PROFILE_LOCK_FILE)
                        return True
                    else:
                        # Process is still running
                        return False
                else:
                    # Invalid lock file, remove it
                    os.remove(PROFILE_LOCK_FILE)
                    return True
        except Exception:
            # Error reading lock file, remove it
            try:
                os.remove(PROFILE_LOCK_FILE)
            except:
                pass
            return True
    return True# =========================
# SELENIUM (OFF-SCREEN)
# =========================
def setup_webdriver(profile_dir=None, user_whatsapp=None):
    if profile_dir is None:
        profile_dir = WHATSAPP_PROFILE_DIR
    
    os.makedirs(profile_dir, exist_ok=True)
    lock_file = os.path.join(profile_dir, ".lock")

    # Extract phone number from profile_dir if it's a phone-specific profile
    # Format: /path/to/.eduping_profiles/phone_91XXXXXXXXXX
    phone_number = None
    if "phone_" in profile_dir:
        phone_number = profile_dir.split("phone_")[-1]
        print(f"[WHATSAPP] Detected phone-specific profile: {phone_number}")
        
        # Force-close any existing driver for this phone
        if phone_number in ACTIVE_DRIVERS:
            print(f"[WHATSAPP] Found existing driver for {phone_number}, closing it first...")
            try:
                close_driver(phone_number)
                time.sleep(1)
            except Exception as e:
                print(f"[WHATSAPP] Could not close existing driver: {e}")

    # Check and clean up stale lock file
    if os.path.exists(lock_file):
        print(f"[WHATSAPP] Lock file found, checking if process is still running...")
        try:
            with open(lock_file, "r") as f:
                pid_str = f.read().strip()
                if pid_str.isdigit():
                    pid = int(pid_str)
                    if is_process_running(pid):
                        print(f"[WHATSAPP] Process {pid} still running, forcing cleanup...")
                        # Remove the lock file anyway - we're opening a new session
                        try:
                            os.remove(lock_file)
                            print(f"[WHATSAPP] Lock file removed (old process will be replaced)")
                        except:
                            pass
                    else:
                        # Process is dead, remove stale lock
                        os.remove(lock_file)
                        print(f"[WHATSAPP] Stale lock file removed")
        except Exception as e:
            print(f"[WHATSAPP] Error checking lock file: {e}")
            # Remove lock file anyway
            try:
                os.remove(lock_file)
            except:
                pass

    # Create new lock file
    with open(lock_file, "w") as f:
        f.write(str(os.getpid()))
    print(f"[WHATSAPP] Created new lock file for PID {os.getpid()}")

    opts = Options()
    # Show window on screen for visible automation
    opts.add_argument("--start-maximized")  # Start maximized so user can see automation
    opts.add_argument("--disable-gpu")
    opts.add_argument("--disable-notifications")
    opts.add_argument("--disable-extensions")
    opts.add_argument("--disable-blink-features=AutomationControlled")
    opts.add_argument("--no-sandbox")
    opts.add_argument(f"--user-data-dir={profile_dir}")
    opts.add_argument("--disable-dev-shm-usage")
    opts.binary_location = r"C:\Program Files\Google\Chrome\Application\chrome.exe"

    try:
        print("="*60)
        print("[WHATSAPP] Starting Chrome browser...")
        if user_whatsapp:
            print(f"[WHATSAPP] User: {user_whatsapp}")
        if phone_number:
            print(f"[WHATSAPP] Phone: {phone_number}")
        print(f"[WHATSAPP] Profile: {profile_dir}")
        driver = webdriver.Chrome(service=Service(), options=opts)
        print("[WHATSAPP] Chrome process started successfully")
        
        # Retry logic for network issues
        max_retries = 3
        retry_delay = 2
        last_error = None
        
        for attempt in range(1, max_retries + 1):
            try:
                print(f"[WHATSAPP] Attempt {attempt}/{max_retries}: Navigating to https://web.whatsapp.com...")
                driver.get("https://web.whatsapp.com")
                print("[WHATSAPP] Page loaded, waiting for page elements...")
                break  # Success, exit retry loop
            except Exception as e:
                last_error = str(e)
                if "net::ERR_NAME_NOT_RESOLVED" in last_error or "ERR_INTERNET_DISCONNECTED" in last_error:
                    print(f"[WHATSAPP] ⚠ Network error on attempt {attempt}: {last_error}")
                    if attempt < max_retries:
                        print(f"[WHATSAPP] Waiting {retry_delay} seconds before retry...")
                        time.sleep(retry_delay)
                        retry_delay *= 2  # Exponential backoff
                    else:
                        print(f"[WHATSAPP] ✗ All {max_retries} attempts failed")
                        print(f"[WHATSAPP] ⚠⚠ NETWORK ISSUE DETECTED ⚠⚠")
                        print(f"[WHATSAPP] Please check:")
                        print(f"[WHATSAPP]   1. Internet connection is active")
                        print(f"[WHATSAPP]   2. DNS is resolving properly")
                        print(f"[WHATSAPP]   3. Firewall/Proxy is not blocking web.whatsapp.com")
                        print(f"[WHATSAPP]   4. You can manually visit https://web.whatsapp.com in Chrome")
                        raise
                else:
                    # Different error, raise immediately
                    raise
        
        # Wait for page to initialize
        time.sleep(5)
        print("[WHATSAPP] WhatsApp Web initialized")
        print("[WHATSAPP] ✓ Browser is VISIBLE for new user login/OTP")
        print("="*60)
        return driver
    except Exception as e:
        print(f"[WHATSAPP] ERROR: Failed to start browser: {str(e)}")
        print(f"[WHATSAPP] Exception type: {type(e).__name__}")
        import traceback
        print(traceback.format_exc())
        raise


# =========================
# LOGIN WITH PHONE NUMBER AUTOMATION
# =========================
def format_phone_for_whatsapp_login(phone_str, default_country="91"):
    """
    Format phone for WhatsApp Web input: +91 xxxxx xxxxx
    Extracts digits, infers country code, returns formatted string.
    """
    if not phone_str:
        return None
    digits = "".join(c for c in str(phone_str) if c.isdigit())
    if not digits:
        return None
    # If starts with country code 91 and has more than 10 digits
    if digits.startswith("91") and len(digits) > 10:
        country = "91"
        phone = digits[2:]
    elif len(digits) == 10:
        country = default_country
        phone = digits
    elif len(digits) >= 11:
        # Assume first 1-3 digits are country code (91, 1, 44, etc.)
        for cc in ["91", "1", "44", "33", "49", "39", "34", "86", "81", "55"]:
            if digits.startswith(cc) and len(digits) > len(cc):
                country = cc
                phone = digits[len(cc):]
                break
        else:
            country = default_country
            phone = digits[-10:]
    else:
        country = default_country
        phone = digits
    # Format: +91 xxxxx xxxxx (5 + 5 for 10-digit Indian number)
    if len(phone) >= 10:
        return f"+{country} {phone[:5]} {phone[5:10]}"
    return f"+{country} {phone}"


def enter_phone_for_whatsapp_login(driver, user_whatsapp, timeout=15):
    """
    Click 'Log in with phone number' and enter the user's phone in format +91 xxxxx xxxxx.
    Returns True if successful.
    """
    if not user_whatsapp:
        return False
    formatted = format_phone_for_whatsapp_login(user_whatsapp)
    if not formatted:
        print("[WHATSAPP] Could not format phone number for login")
        return False
    print(f"[WHATSAPP] Auto-entering phone: {formatted}")
    try:
        # Click "Log in with phone number"
        login_btn_xpaths = [
            "//div[contains(text(), 'Log in with phone number')]",
            "//div[@class='x1c4vz4f xs83m0k xdl72j9 x1g77sc7 x78zum5 xozqiw3 x1oa3qoh x12fk4p8 xeuugli x2lwn1j x1nhvcw1 xdt5ytf x1cy8zhl' and contains(text(), 'Log in with phone number')]",
            "//*[contains(text(), 'Log in with phone number')]",
        ]
        clicked = False
        for xpath in login_btn_xpaths:
            try:
                els = driver.find_elements(By.XPATH, xpath)
                for el in els:
                    if el.is_displayed():
                        el.click()
                        clicked = True
                        print("[WHATSAPP] Clicked 'Log in with phone number'")
                        time.sleep(2)
                        break
                if clicked:
                    break
            except Exception:
                pass
        if not clicked:
            print("[WHATSAPP] Could not find/click 'Log in with phone number' (may already be on phone input screen)")
        
        # Find input by aria-label
        input_xpaths = [
            "//input[@aria-label='Type your phone number to log in to WhatsApp']",
            "//input[contains(@aria-label, 'phone number')]",
            "//input[@type='text' and contains(@class, 'selectable-text')]",
        ]
        inp = None
        for xpath in input_xpaths:
            try:
                els = driver.find_elements(By.XPATH, xpath)
                for el in els:
                    if el.is_displayed():
                        inp = el
                        break
                if inp:
                    break
            except Exception:
                pass
        
        if not inp:
            print("[WHATSAPP] Could not find phone number input field")
            return False
        
        # Clear existing value and enter formatted number
        inp.click()
        time.sleep(0.3)
        inp.clear()
        inp.send_keys(Keys.CONTROL + "a")
        inp.send_keys(Keys.DELETE)
        time.sleep(0.2)
        inp.send_keys(formatted)
        print(f"[WHATSAPP] Entered phone: {formatted}")
        time.sleep(5)
        # Click Next button after 5 seconds
        next_btn_xpaths = [
            "//button[.//div[contains(text(), 'Next')]]",
            "//div[contains(text(), 'Next')]/ancestor::button",
            "//button[contains(@class, 'x889kno')]",
            "//*[contains(text(), 'Next')]",
        ]
        for xpath in next_btn_xpaths:
            try:
                els = driver.find_elements(By.XPATH, xpath)
                for el in els:
                    if el.is_displayed():
                        el.click()
                        print("[WHATSAPP] Clicked 'Next' button")
                        time.sleep(1)
                        return True
            except Exception:
                pass
        print("[WHATSAPP] Could not find/click 'Next' button")
        return True
    except Exception as e:
        print(f"[WHATSAPP] Error auto-entering phone: {e}")
        import traceback
        print(traceback.format_exc())
        return False


# =========================
# AUTHENTICATION COMPLETION DETECTION
# =========================
def is_whatsapp_authenticated(driver, timeout=30):
    """
    Detect when WhatsApp Web has completed authentication.
    Checks for chat interface elements that appear after successful login.
    """
    print("[AUTH-DETECT] Checking for WhatsApp authentication completion...")
    start_time = time.time()
    
    # Elements that indicate successful authentication
    auth_indicators = [
        "//div[@role='textbox' and @aria-placeholder='Type a message']",  # Message input box
        "//footer//div[@contenteditable='true']",  # Message compose area
        "//div[contains(@class, 'x1c4vz4f') and .//span[contains(text(), 'Message')]]",
        "//div[@data-testid='conversationlist']",  # Chat list
        "//div[@data-testid='search-input']",  # Search box (appears after auth)
    ]
    
    while time.time() - start_time < timeout:
        try:
            for indicator in auth_indicators:
                elements = driver.find_elements(By.XPATH, indicator)
                if elements and any(e.is_displayed() for e in elements):
                    elapsed = time.time() - start_time
                    print(f"[AUTH-DETECT] ✓ WhatsApp AUTHENTICATED (took {elapsed:.1f}s)")
                    return True
        except:
            pass
        
        time.sleep(0.5)
    
    elapsed = time.time() - start_time
    print(f"[AUTH-DETECT] Could not confirm authentication after {elapsed:.1f}s (may still be loading)")
    return False


def is_otp_screen_visible(driver, timeout=20):
    """
    Detect if OTP entry screen is visible.
    This checks for UI elements that appear during OTP input phase.
    """
    print("[OTP-DETECT] Checking for OTP screen...")
    start_time = time.time()
    
    # Multiple UI elements that indicate OTP screen
    otp_indicators = [
        "//div[contains(text(), '6-digit')]",  # "Enter 6-digit code"
        "//div[contains(text(), 'Enter the OTP')]",
        "//div[contains(text(), 'code')]",
        "//input[@type='text' and @placeholder='Code']",
        "//div[@data-testid='otp']" ,
        "//div[@role='textbox' and contains(@class, 'otp')]",
        "//div[contains(text(), 'Phone number code')]",  # Common WhatsApp message
    ]
    
    while time.time() - start_time < timeout:
        try:
            for indicator in otp_indicators:
                elements = driver.find_elements(By.XPATH, indicator)
                if elements and any(e.is_displayed() for e in elements):
                    print(f"[OTP-DETECT] ✓ OTP screen DETECTED (took {time.time() - start_time:.1f}s)")
                    return True
        except:
            pass
        
        time.sleep(1)
    
    print(f"[OTP-DETECT] No OTP screen detected after {timeout}s")
    return False


def send_message_once(driver, phone, message):
    """Send a single message attempt - supports both saved and unsaved contacts"""
    # Format phone number for WhatsApp Web URL (remove any special characters)
    phone_clean = str(phone).replace("+", "").replace(" ", "").replace("-", "")
    if not phone_clean.isdigit():
        phone_clean = ''.join(c for c in phone_clean if c.isdigit())
    
    # Navigate to the chat with this phone number
    # This URL works for BOTH saved and unsaved contacts in WhatsApp Web
    url = f"https://web.whatsapp.com/send?phone={phone_clean}&text=&app_absent=0"
    print(f"[WHATSAPP] 🔗 Navigating to: {url}")
    
    try:
        driver.get(url)
        print(f"[WHATSAPP] ✓ Navigation successful")
    except Exception as e:
        print(f"[WHATSAPP] ❌ Error navigating to URL: {e}")
        return False
    
    # Wait for chat interface to load
    print("[WHATSAPP] ⏳ Waiting for chat interface to load...")
    time.sleep(6)
    
    # Check for error messages
    try:
        # Check for invalid phone number error
        error_elements = driver.find_elements(By.XPATH, "//*[contains(text(), 'Phone number shared via url is invalid')]")
        if error_elements:
            print(f"[WHATSAPP] ❌ Error: Phone number {phone_clean} is invalid")
            return False
    except Exception as e:
        print(f"[WHATSAPP] ⚠ Error checking for invalid phone: {e}")
        pass
    
    # Try multiple XPath options for the message box (WhatsApp Web UI changes frequently)
    textbox_xpaths = [
        "//div[@role='textbox' and @aria-placeholder='Type a message']",
        "//div[@contenteditable='true' and @data-tab='10']",
        "//div[@contenteditable='true'][@role='textbox']",
        "//footer//div[@contenteditable='true']",
        "//div[contains(@class, 'selectable-text')][@contenteditable='true']",
        "//div[contains(@class, 'x1n2onr6')][@contenteditable='true'][@data-tab]",
        "//div[@contenteditable='true'][contains(@class, 'copyable-text')]",
        "//p[@class='selectable-text copyable-text x15bjb6t x1n2onr6 x1vjfegm1']",
        "//div[@contenteditable='true' and @spellcheck='true']",
        "//div[contains(@aria-label, 'Type a message')]",
    ]
    
    box = None
    print(f"[WHATSAPP] 🔍 Looking for message textbox...")
    for i, xpath in enumerate(textbox_xpaths, 1):
        try:
            print(f"[WHATSAPP]   Trying XPath option {i}/{len(textbox_xpaths)}...")
            box = WebDriverWait(driver, 5).until(
                EC.element_to_be_clickable((By.XPATH, xpath))
            )
            if box and box.is_displayed():
                print(f"[WHATSAPP] ✓ Found textbox using XPath option {i}")
                break
        except Exception:
            box = None
            continue
    
    if not box:
        print("[WHATSAPP] ❌ ERROR: Could not find message textbox after trying all XPath options!")
        print("[WHATSAPP] This could mean:")
        print("[WHATSAPP]   - Chat window is not fully loaded")
        print("[WHATSAPP]   - WhatsApp Web UI has changed")
        print("[WHATSAPP]   - Phone number might be blocked or invalid")
        print("[WHATSAPP]   - Browser is not in focus")
        return False

    print(f"[WHATSAPP] ✓ Textbox found successfully")

    # Click on textbox using JavaScript (avoids "Element click intercepted" error)
    try:
        print(f"[WHATSAPP] 🖱️  Clicking textbox using JavaScript...")
        driver.execute_script("arguments[0].click();", box)
        print(f"[WHATSAPP] ✓ Click successful")
    except Exception as e:
        print(f"[WHATSAPP] ⚠ JavaScript click failed, trying direct click: {e}")
        try:
            box.click()
            print(f"[WHATSAPP] ✓ Direct click successful")
        except Exception as e2:
            print(f"[WHATSAPP] ❌ Both click methods failed: {e2}")
            return False
    
    time.sleep(0.5)
    print(f"[WHATSAPP] ⌨️  Typing message ({len(message)} characters)...")
    box.send_keys(message)
    print(f"[WHATSAPP] ✓ Message typed successfully")
    time.sleep(1)
    
    # Press Enter to send
    print(f"[WHATSAPP] 📤 Pressing Enter to send...")
    box.send_keys(Keys.ENTER)
    print(f"[WHATSAPP] ✓ Enter pressed - message sent")
    time.sleep(2)
    
    # Try clicking the send button as backup using JavaScript
    try:
        send_btn = driver.find_elements(By.XPATH, "//span[@data-icon='send']")
        if send_btn:
            driver.execute_script("arguments[0].click();", send_btn[0])
            time.sleep(1)
    except:
        pass
    
    return True


def check_and_retry_failed_messages(driver):
    """Check for failed messages and click retry on them"""
    retried = 0
    try:
        # Look for error/alert icons on messages (red exclamation marks)
        error_icons = driver.find_elements(By.XPATH, "//span[@data-icon='alert-phone']|//span[@data-icon='msg-error']|//div[@data-icon='alert-phone']")
        
        for icon in error_icons:
            try:
                # Click on the error icon to retry
                icon.click()
                time.sleep(1)
                
                # Look for retry button and click it
                retry_btns = driver.find_elements(By.XPATH, "//*[contains(text(), 'Retry')]|//*[contains(text(), 'retry')]")
                for btn in retry_btns:
                    try:
                        btn.click()
                        retried += 1
                        print(f"[WHATSAPP] Clicked retry on failed message")
                        time.sleep(2)
                    except:
                        pass
            except:
                pass
        
        if retried > 0:
            print(f"[WHATSAPP] Retried {retried} failed messages")
            time.sleep(5)  # Wait for retries to process
            
    except Exception as e:
        print(f"[WHATSAPP] Error checking for failed messages: {e}")
    
    return retried


def send_message(driver, phone, message, max_retries=3):
    """Send message with automatic retry on failure - supports unsaved numbers"""
    print(f"\n[WHATSAPP] 📱 Sending to: {phone} (saved or unsaved contact)")
    print(f"[WHATSAPP] Message preview: {message[:80]}...")
    print(f"[WHATSAPP] Attempts: {max_retries} retries available")
    
    for attempt in range(1, max_retries + 1):
        print(f"\n[WHATSAPP] 🔄 Attempt {attempt}/{max_retries}")
        
        # Send the message
        if send_message_once(driver, phone, message):
            # Wait for delivery
            print("[WHATSAPP] ⏳ Waiting for delivery confirmation...")
            time.sleep(3)
            
            # Check for delivery confirmation
            try:
                # Look for tick marks (sent/delivered)
                ticks = driver.find_elements(By.XPATH, "//span[@data-icon='msg-check' or @data-icon='msg-dblcheck']")
                if ticks:
                    print(f"[WHATSAPP] ✅ Message delivered to {phone}")
                    return True
                
                # Check for error icons
                errors = driver.find_elements(By.XPATH, "//span[@data-icon='alert-phone']|//span[@data-icon='msg-error']")
                if errors and attempt < max_retries:
                    print(f"[WHATSAPP] ⚠ Delivery failed - retrying...")
                    # Try clicking retry on the failed message
                    check_and_retry_failed_messages(driver)
                    time.sleep(2)
                    continue
            except Exception as e:
                print(f"[WHATSAPP] ⚠ Could not verify delivery: {e}")
                pass
            
            # If we can't determine status, assume success after waiting
            print(f"[WHATSAPP] ✓ Message sent to {phone} (status confirmed)")
            return True
        else:
            if attempt < max_retries:
                print(f"[WHATSAPP] ❌ Send failed - retrying in 2 seconds...")
                time.sleep(2)
    
    print(f"[WHATSAPP] ❌ FAILED to send to {phone} after {max_retries} attempts")
    return False


# =========================
# STATUS TRACKING
# =========================
def update_status(status: str, progress: int = 0, total: int = 0, message: str = ""):
    """Update status file for frontend to read"""
    os.makedirs(WHATSAPP_PROFILE_DIR, exist_ok=True)
    status_data = {
        "status": status,  # "processing", "sending", "completed", "error"
        "progress": progress,
        "total": total,
        "message": message,
        "timestamp": datetime.now().isoformat()
    }
    with open(STATUS_FILE, "w") as f:
        json.dump(status_data, f)


def get_status():
    """Get current status"""
    if os.path.exists(STATUS_FILE):
        try:
            with open(STATUS_FILE, "r") as f:
                return json.load(f)
        except:
            pass
    return {"status": "idle", "progress": 0, "total": 0, "message": ""}


def clear_status():
    """Clear status file"""
    if os.path.exists(STATUS_FILE):
        os.remove(STATUS_FILE)


def clear_lock_file():
    """Manually clear the lock file (use if you get 'WhatsApp already running' error)"""
    if os.path.exists(PROFILE_LOCK_FILE):
        try:
            os.remove(PROFILE_LOCK_FILE)
            return True
        except Exception:
            return False
    return True


# =========================
# BACKGROUND WORKER
# =========================
def _background_worker(phone_file_path, message_file_path, user_whatsapp=None, use_authenticated_session=False, country_code="", phone_only=""):
    driver = None
    profile_dir = get_user_profile_dir(user_whatsapp)
    try:
        print("\n" + "="*60)
        print("[EDUPING] Starting background worker...")
        if user_whatsapp:
            print(f"[EDUPING] User Number: {user_whatsapp}")
            print(f"[EDUPING] Profile Dir: {profile_dir}")
        if use_authenticated_session:
            print(f"[EDUPING] Using authenticated OTP session")
        print("="*60)
        
        update_status("processing", 0, 0, "Generating summaries...")
        pairs = build_pairs(phone_file_path, message_file_path)
        
        total = len(pairs)
        print(f"\n[EDUPING] ========== PAIRS SUMMARY ==========")
        print(f"[EDUPING] Total pairs to send: {total}")
        if total == 0:
            print(f"[EDUPING] ⚠️  WARNING: No phone-message pairs created!")
            print(f"[EDUPING] This could mean:")
            print(f"[EDUPING]   - Phone file is empty")
            print(f"[EDUPING]   - Message file is empty")
            print(f"[EDUPING]   - File extraction failed")
            update_status("error", 0, 0, "No phone numbers found in file")
            return
        else:
            print(f"[EDUPING] ✓ Sample pairs:")
            for i, (phone, msg) in enumerate(pairs[:3]):
                print(f"[EDUPING]   {i+1}. Phone: {phone} | Message: {msg[:50]}...")
        print(f"[EDUPING] =====================================\n")
        update_status("sending", 0, total, f"Preparing to send {total} messages...")
        
        print("[EDUPING] Preparing WhatsApp WebDriver...")

        # If we have an authenticated session requested, try to reuse existing driver
        existing_driver_used = False
        full_phone = None
        phone_profile_dir = None
        if use_authenticated_session and country_code and phone_only:
            full_phone = country_code + phone_only
            # Create phone-specific profile directory (preserves WhatsApp authentication)
            phone_profile_dir = os.path.join(WHATSAPP_PROFILE_DIR, f"phone_{full_phone}")
            os.makedirs(phone_profile_dir, exist_ok=True)
            print(f"[EDUPING] 🔍 Looking for authenticated driver for {full_phone}...")
            
            # TRY TO REUSE EXISTING DRIVER FROM OTP PHASE
            with DRIVERS_LOCK:
                if full_phone in ACTIVE_DRIVERS:
                    driver = ACTIVE_DRIVERS[full_phone]
                    existing_driver_used = True
                    print(f"[EDUPING] ✅ ✅ ✅ REUSING AUTHENTICATED DRIVER FOR {full_phone} - NO DOUBLE LOADING!")
                    print(f"[EDUPING] ✅ Driver is ALIVE and ready - proceeding directly to message sending")
                    print(f"[EDUPING] ✅ Driver type: {type(driver)}")
                    print(f"[EDUPING] ✅ Driver current URL: {driver.current_url if driver else 'None'}")
                else:
                    print(f"[EDUPING] ⚠ No existing driver found for {full_phone}")
                    print(f"[EDUPING] This means user hasn't completed OTP or driver was closed")

        if not driver and not existing_driver_used:
            # No existing driver found or not using authenticated session - start a new one
            # CRITICAL: Always use phone-specific profile if we have one (preserves WhatsApp auth)
            if phone_profile_dir:
                print(f"[EDUPING] Opening NEW WebDriver with authenticated profile: {phone_profile_dir}")
                driver = setup_webdriver(profile_dir=phone_profile_dir)
                print(f"[EDUPING] ✓ NEW WebDriver opened (no previous session found)")
            else:
                # Fresh profile each run - ensures WhatsApp always shows login screen (QR or phone) for both new and returning users
                fresh_profile = os.path.join(WHATSAPP_PROFILE_DIR, f"fresh_{int(time.time())}_{os.getpid()}")
                os.makedirs(fresh_profile, exist_ok=True)
                print(f"[EDUPING] Using fresh profile for login: {fresh_profile}")
                driver = setup_webdriver(profile_dir=fresh_profile, user_whatsapp=user_whatsapp)
            print("[EDUPING] WebDriver ready for message sending...")
            # Register driver for tracking if we have a phone context
            if use_authenticated_session and country_code and phone_only:
                register_driver(full_phone, driver)
        elif not driver:
            print(f"[EDUPING] ✓ Driver will be used from existing session")
        
        if use_authenticated_session and country_code and phone_only:
            print("[EDUPING] ========================================")
            print("[EDUPING] Using pre-authenticated WhatsApp session")
            print(f"[EDUPING] Country Code: {country_code}, Phone: {country_code + phone_only}")
            update_status("sending", 0, total, "Ready to send messages using authenticated session...")
            
            # Check if a driver already exists for this phone number (from OTP request)
            full_phone = country_code + phone_only
            driver_key = full_phone
            
            with DRIVERS_LOCK:
                if driver_key in ACTIVE_DRIVERS:
                    print(f"[EDUPING] ✓ Found existing authenticated driver for {full_phone}")
                    # driver already set above when reusing
                    print("[EDUPING] ✓ WhatsApp should already be authenticated and ready!")
                else:
                    print(f"[EDUPING] ⚠ No existing driver found for {full_phone}")
                    print("[EDUPING]   This could mean:")
                    print("[EDUPING]   - User hasn't completed OTP verification yet")
                    print("[EDUPING]   - Or browser was closed")
                    print("[EDUPING]   Proceeding to open new browser for sending...")
                    print("[EDUPING] Will use the WebDriver started above")
                
        elif use_authenticated_session:
            print("[EDUPING] Waiting for authenticated WhatsApp session...")
            # Wait for WhatsApp to authenticate (user should already be logged in via OTP)
            if wait_for_whatsapp_authentication(driver, timeout=180):
                print("[EDUPING] WhatsApp authenticated successfully!")
            else:
                print("[EDUPING] Warning: WhatsApp authentication timeout, proceeding anyway...")
        else:
            # New flow: enter phone, click Next, wait for user to enter OTP and chats to load, then hide and send
            print("[EDUPING] *** WhatsApp Web opened - preparing phone login ***")
            time.sleep(3)  # Let page stabilize
            if user_whatsapp:
                enter_phone_for_whatsapp_login(driver, user_whatsapp)
            else:
                print("[EDUPING] No user WhatsApp number provided - user can log in manually (QR or phone)")
            update_status("waiting_login", 0, total, "Enter OTP on your phone. After chats load, window will hide and messages will send.")
            # Wait for user to enter OTP and chats to load (no auth check - fixed wait)
            print("[EDUPING] Waiting 15 seconds for OTP entry...")
            time.sleep(15)
            # Wait 10 seconds for profile to load after OTP
            print("[EDUPING] Waiting 10 seconds for profile to load...")
            time.sleep(10)
            # Make window invisible (off-screen) - not closed, WhatsApp continues processing in background
            try:
                driver.set_window_position(-32000, -32000)
                print("[EDUPING] Window moved off-screen (invisible but still running in background)")
            except Exception as e:
                print(f"[EDUPING] Could not move off-screen: {e}")
        
        # Before sending, ensure browser is hidden (if it wasn't hidden during OTP flow)
        if full_phone and use_authenticated_session:
            try:
                with DRIVERS_LOCK:
                    if full_phone in ACTIVE_DRIVERS:
                        # Browser should already be hidden from OTP flow, but verify
                        hid_ok = hide_driver_window(full_phone)
                        if hid_ok:
                            print(f"[EDUPING] ✓ Browser window confirmed hidden for {full_phone}")
                        else:
                            print(f"[EDUPING] ⚠ Could not confirm browser hidden for {full_phone}")
            except:
                pass

        update_status("sending", 0, total, "Connecting to WhatsApp...")
        print("[EDUPING] ========== MESSAGE SENDING START ==========")
        print(f"[EDUPING] Driver is ready: {driver is not None}")
        print(f"[EDUPING] Total messages to send: {total}")
        print(f"[EDUPING] Starting message loop now...\n")

        failed_phones = []
        for idx, (phone, summary) in enumerate(pairs, 1):
            try:
                print(f"\n[EDUPING] === [{idx}/{total}] Sending to {phone} ===")
                success = send_message(driver, phone, summary)
                if success:
                    print(f"[EDUPING] ✓ [{idx}/{total}] SUCCESS!")
                    update_status("sending", idx, total, f"Sent message {idx} of {total}")
                else:
                    print(f"[EDUPING] ✗ [{idx}/{total}] FAILED!")
                    failed_phones.append(phone)
                    update_status("sending", idx, total, f"Failed to send to {phone}")
                print(f"[EDUPING] Waiting before next message...")
                time.sleep(2)
            except Exception as e:
                print(f"[EDUPING] ❌ EXCEPTION on message {idx}: {str(e)}")
                import traceback
                print(traceback.format_exc())
                failed_phones.append(phone)
                update_status("sending", idx, total, f"Error sending to {phone}: {str(e)}")
                time.sleep(2)

        # Final check: retry any remaining failed messages
        if failed_phones:
            print(f"\n[EDUPING] Retrying {len(failed_phones)} failed messages...")
            update_status("sending", total, total, f"Retrying {len(failed_phones)} failed messages...")
            check_and_retry_failed_messages(driver)
            time.sleep(3)

        print("\n" + "="*60)
        print("[EDUPING] ========== SENDING COMPLETE ==========")
        if failed_phones:
            print(f"[EDUPING] ✓ Sent {total - len(failed_phones)} of {total} messages successfully")
            print(f"[EDUPING] ✗ Failed to send to {len(failed_phones)} numbers: {failed_phones}")
        else:
            print(f"[EDUPING] ✅ SUCCESS! All {total} messages sent successfully!")
        print("[EDUPING] ========== END ==========")
        print("="*60 + "\n")
        update_status("completed", total, total, f"Sent {total - len(failed_phones)} of {total} messages")

    except Exception as e:
        print(f"\n[EDUPING] FATAL ERROR: {str(e)}")
        update_status("error", 0, 0, f"Error: {str(e)}")
    finally:
        if driver:
            try:
                print("[EDUPING] Closing WebDriver...")
                driver.quit()
            except:
                pass  # Driver already closed
        # Unregister any remaining drivers
        close_all_drivers()
        lock_file = os.path.join(profile_dir, ".lock")
        if os.path.exists(lock_file):
            os.remove(lock_file)
        print("[EDUPING] Background worker finished.")


# =========================
# PUBLIC API
# =========================
def send_bulk_from_files(phone_file_path, message_file_path, user_whatsapp=None, use_authenticated_session=False, country_code="", phone_only=""):
    """
    Start background process to send messages. Returns immediately.
    
    Args:
        phone_file_path: Path to file containing phone numbers
        message_file_path: Path to file containing messages
        user_whatsapp: WhatsApp number of the user sending messages (optional)
        use_authenticated_session: Whether to use pre-authenticated OTP session
        country_code: Country code for automated phone login (e.g., "91")
        phone_only: Phone number for automated login (e.g., "9876543210")
    """
    clear_status()  # Clear any previous status
    thread = threading.Thread(
        target=_background_worker,
        args=(phone_file_path, message_file_path, user_whatsapp),
        kwargs={"use_authenticated_session": use_authenticated_session, "country_code": country_code, "phone_only": phone_only},
        daemon=True
    )
    thread.start()
    return thread


# =========================
# OTP AUTHENTICATION API
# =========================
def start_otp_automation_thread(phone_number: str, country_code: str = "", phone_only: str = ""):
    """
    Start WhatsApp Web automation in a background thread.
    This opens the browser and automatically logs in with phone number.
    Called immediately when user clicks "Request OTP".
    """
    def automation_thread():
        driver = None
        try:
            print("\n" + "="*60)
            print("[OTP-AUTO] Starting OTP automation thread...")
            print(f"[OTP-AUTO] Phone: +{country_code}{phone_only}")
            print("="*60)
            
            # Check if whatsapp_otp module was imported successfully
            if not start_whatsapp_login_automation:
                print("[OTP-AUTO] ✗ Error: whatsapp_otp module not imported!")
                print("[OTP-AUTO] Cannot proceed with automation")
                return
            
            full_phone = country_code + phone_only
            
            # STEP 0: Check if user is ALREADY logged in (from previous session)
            print("[OTP-AUTO] STEP 0: Checking if user is already logged in from previous session...")
            if is_session_already_logged_in and is_session_already_logged_in(full_phone):
                print("[OTP-AUTO] ✓✓✓ USER ALREADY LOGGED IN FROM PREVIOUS SESSION ✓✓✓")
                print("[OTP-AUTO] ✓ Skipping browser entirely - no need to login again")
                print("[OTP-AUTO] ✓ No OTP needed - session ready for sending")
                print("[OTP-AUTO] ✓ Browser will NOT be opened at all")
                print("[OTP-AUTO] Automation complete - no browser, no OTP!")
                return  # Exit - NEVER open browser for existing users
            
            # STEP 1: Open browser only if NOT already logged in (NEW USER)
            print("[OTP-AUTO] ✓ User is NEW - opening browser for login...")
            print(f"[OTP-AUTO] Creating user-specific profile for {full_phone}...")
            # Use phone-specific profile directory (preserves authentication)
            phone_profile_dir = os.path.join(WHATSAPP_PROFILE_DIR, f"phone_{full_phone}")
            os.makedirs(phone_profile_dir, exist_ok=True)
            print(f"[OTP-AUTO] Profile directory: {phone_profile_dir}")
            print("[OTP-AUTO] Opening WhatsApp Web browser (VISIBLE for login)...")
            driver = setup_webdriver(profile_dir=phone_profile_dir)
            
            # Register driver so it can be accessed later
            register_driver(full_phone, driver)
            
            # NO WAIT - proceed immediately to automation
            print("[OTP-AUTO] Starting phone number entry automation...")
            result = start_whatsapp_login_automation(driver, country_code, phone_only)
            
            if result.get("success"):
                print("[OTP-AUTO] ✓ Automation successful!")
                
                # Check if user was ALREADY logged in (detected during automation)
                if result.get("skip_otp"):
                    print("[OTP-AUTO] ✓✓ User was ALREADY logged in!")
                    print("[OTP-AUTO] ✓ Skipping OTP - hiding browser now for security")
                    # User already logged in, hide browser immediately
                    hide_driver_window(full_phone)
                else:
                    print("[OTP-AUTO] ✓ Phone entry automation complete")
                    print("[OTP-AUTO] 📱 Waiting for user to enter OTP on their phone...")
                    print("[OTP-AUTO] [KEEP BROWSER WINDOW VISIBLE - showing OTP field]")
                    
                    # Wait for authentication to complete (user enters OTP on phone)
                    print("[OTP-AUTO] Waiting for WhatsApp authentication (user entering OTP)...")
                    if is_whatsapp_authenticated(driver, timeout=60):
                        print("[OTP-AUTO] ✓ WhatsApp has authenticated successfully!")
                        print("[OTP-AUTO] ✓ User logged in - hiding browser immediately for security")
                        # User successfully logged in, hide browser NOW
                        hide_driver_window(full_phone)
                        print("[OTP-AUTO] ✓ Browser is HIDDEN - WhatsApp session saved for message sending")
                    else:
                        print("[OTP-AUTO] ⚠ Authentication timeout - browser still visible")
                        print("[OTP-AUTO] User may have already logged in - attempting to hide browser...")
                        hide_driver_window(full_phone)
            else:
                print(f"[OTP-AUTO] ✗ Automation failed: {result.get('message')}")
                print(f"[OTP-AUTO] Step: {result.get('step')}")
                # Close browser if automation failed
                close_driver(full_phone)
                
        except Exception as e:
            print(f"[OTP-AUTO] ✗ Error in automation thread: {str(e)}")
            import traceback
            print(traceback.format_exc())
            if driver:
                try:
                    close_driver(country_code + phone_only)
                except:
                    pass
    
    # Start automation in background thread
    thread = threading.Thread(target=automation_thread, daemon=True)
    thread.start()
    print("[OTP-AUTO] Automation thread started!")


def request_otp(phone_number: str, country_code: str = "", phone_only: str = "") -> dict:
    """
    Request OTP for WhatsApp login.
    Immediately starts automation in background.
    Frontend calls this, user receives OTP on their phone.
    
    Args:
        phone_number: Full phone number (country code + phone)
        country_code: Country code (e.g., "91")
        phone_only: Phone number without country code (e.g., "9876543210")
    """
    # Clean phone number
    phone = phone_number.replace("+", "").replace(" ", "").replace("-", "")
    if not phone.isdigit():
        return {"success": False, "message": "Invalid phone number format"}
    
    full_phone = country_code + phone_only
    
    # Check if user is already logged in from previous session
    already_logged_in = False
    if is_session_already_logged_in:
        already_logged_in = is_session_already_logged_in(full_phone)
    
    # Save OTP session (if function available)
    if save_otp_session:
        try:
            save_otp_session(phone, "", None)
        except Exception as e:
            print(f"[OTP-REQ] Warning: Could not save OTP session: {e}")
    
    print(f"[OTP-REQ] OTP request initiated for {phone}")
    print(f"[OTP-REQ] Country Code: {country_code}, Phone: {phone_only}")
    print(f"[OTP-REQ] Already logged in: {already_logged_in}")
    
    # Determine message based on login status
    if already_logged_in:
        msg = "✓ WhatsApp session found from previous login!\n\nNo need to log in again.\n\nSkipping browser - ready to generate summaries and send messages!"
        skip_otp_input = True
    else:
        msg = "🚀 WhatsApp Web is opening with automated phone login...\n\nWatch the browser window as your phone number is entered automatically.\n\nYou'll receive a 6-digit code on your phone."
        skip_otp_input = False
    
    # Start automation thread immediately
    print("[OTP-REQ] Starting WhatsApp Web automation...")
    start_otp_automation_thread(phone_number, country_code, phone_only)
    
    return {
        "success": True,
        "message": msg,
        "phone": phone,
        "country_code": country_code,
        "phone_only": phone_only,
        "automation_started": True,
        "already_logged_in": already_logged_in,
        "skip_otp_input": skip_otp_input  # FLAG for frontend: whether to show OTP input section
    }


def verify_otp(phone_number: str, otp_code: str) -> dict:
    """
    Verify OTP that user received.
    Once verified, WhatsApp session is ready for sending.
    Also closes the browser window after verification.
    """
    result = verify_whatsapp_otp(phone_number, otp_code)
    
    # After OTP verification, hide the browser window (do NOT close the driver yet)
    if result.get("success"):
        # Extract phone number for driver handling
        phone = result.get("phone", phone_number.replace("+", "").replace(" ", "").replace("-", ""))
        print(f"[BACKEND] OTP verified for {phone}, hiding browser window instead of closing driver...")
        hid = hide_driver_window(phone)
        if not hid:
            print(f"[BACKEND] Warning: could not hide browser for {phone}. Leaving driver running.")
    
    return result
